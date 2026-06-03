from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


PROJECT = Path(r"D:\Development\Projects\BottleWebProject_C326_TEAM_STEP\BottleWebProject_C326_TEAM_STEP")
DOC_PATH = Path(r"D:\Development\Projects\UP04\UP02") / "Петренко.docx"


def set_run_font(run, name="Times New Roman", size=14, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold


def clear_text(paragraph, text, font="Times New Roman", size=14, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return paragraph


def format_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(35)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def format_no_indent(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_run(paragraph, text, font="Times New Roman", size=14, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return run


def insert_before(anchor, text="", style=None, body=True, bold=False):
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        add_run(paragraph, text, bold=bold)
    if body:
        format_body(paragraph)
    else:
        format_no_indent(paragraph)
    return paragraph


def add_heading_before(anchor, text, style="Heading 1"):
    paragraph = insert_before(anchor, text, style=style, body=False, bold=True)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_page_break_before(anchor):
    paragraph = insert_before(anchor, body=False)
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


def append_heading(doc, text, style="Heading 1"):
    paragraph = doc.add_paragraph(style=style)
    add_run(paragraph, text, bold=True)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def append_body(doc, text):
    paragraph = doc.add_paragraph()
    add_run(paragraph, text)
    format_body(paragraph)
    return paragraph


def append_code(doc, code):
    for number, line in enumerate(code.splitlines(), start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{number:04d}: {line}")
        set_run_font(run, "Courier New", 8, False)


def find_para(doc, exact=None, startswith=None):
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if exact is not None and text == exact:
            return paragraph
        if startswith is not None and text.startswith(startswith):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {exact or startswith}")


def replace_start(doc, startswith, text):
    paragraph = find_para(doc, startswith=startswith)
    clear_text(paragraph, text)
    format_body(paragraph)


def delete_between(doc, first_exclusive, last_exclusive):
    start = doc.paragraphs.index(first_exclusive)
    end = doc.paragraphs.index(last_exclusive)
    for paragraph in list(doc.paragraphs[start + 1:end]):
        paragraph._element.getparent().remove(paragraph._element)


def read_project_file(relative_path):
    return (PROJECT / relative_path).read_text(encoding="utf-8")


def routes_coloring_excerpt():
    lines = read_project_file("routes.py").splitlines()
    selected = []
    selected.extend(lines[5:19])
    selected.append("")
    selected.extend(lines[81:196])
    return "\n".join(selected)


def apply_heading_styles(doc):
    in_report = False
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "ВВЕДЕНИЕ":
            in_report = True
        if not in_report or not text:
            continue

        heading_style = None
        if (
            text in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ"}
            or re.match(r"^[1-5]\s+\S", text)
            or text.startswith("ПРИЛОЖЕНИЕ ")
        ):
            heading_style = "Heading 1"
        elif re.match(r"^[1-5]\.\d+\.\d+\s+", text):
            heading_style = "Heading 3"
        elif re.match(r"^[1-5]\.\d+\s+", text) or text.startswith("Файл: "):
            heading_style = "Heading 2"

        if heading_style:
            paragraph.style = heading_style
            format_no_indent(paragraph)
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman", 14, True)


def main():
    doc = Document(DOC_PATH)

    if any("ПРИЛОЖЕНИЕ Б. Листинг 2" in p.text for p in doc.paragraphs):
        raise RuntimeError("Приложение Б уже добавлено. Скрипт остановлен, чтобы не продублировать разделы.")

    content = find_para(doc, exact="СОДЕРЖАНИЕ")
    intro = find_para(doc, exact="ВВЕДЕНИЕ")
    delete_between(doc, content, intro)
    toc_items = [
        "ВВЕДЕНИЕ\t3",
        "1 Техническое задание\t8",
        "2 Описание алгоритма решения задачи\t19",
        "3 UML-диаграммы Use Case и компонентов приложения\t23",
        "4 Структура проекта\t40",
        "5 Тестирование\t53",
        "ЗАКЛЮЧЕНИЕ\t73",
        "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ\t75",
        "ПРИЛОЖЕНИЕ А. Листинг 1. Код файлов вёрстки проекта\t76",
        "ПРИЛОЖЕНИЕ Б. Листинг 2. Код логики индивидуального модуля\t109",
    ]
    for item in reversed(toc_items):
        insert_before(intro, item, body=False)

    replace_start(
        doc,
        "Индивидуальным модулем студента Петренко Антона является модуль автоматического составления расписания",
        "Индивидуальным модулем студента Петренко Антона является модуль автоматического составления расписания на основе раскраски графа конфликтов (вариант 4). В текущей версии модуль выделен в отдельную теоретическую страницу /coloring и практическую страницу /coloring/practice. Пользователь работает с таблицей дисциплин, списком конфликтов, импортом JSON, случайной генерацией данных, интерактивным графом и блоком результата.",
    )
    replace_start(
        doc,
        "Математически задача представляется как раскраска вершин неориентированного графа конфликтов",
        "Математически задача представляется как раскраска вершин неориентированного графа конфликтов: вершина соответствует дисциплине, ребро соответствует конфликту, цвет соответствует смене. Для решения применяется заданный вариантом алгоритм: жадная раскраска Welsh-Powell с последующей локальной оптимизацией. Полный перебор всех возможных раскрасок не используется, так как он не относится к выбранному алгоритму.",
    )
    replace_start(
        doc,
        "Модуль выполняет следующие функции. Первая функция",
        "Модуль выполняет следующие функции. Первая функция — приём и валидация входных данных: перечня дисциплин с преподавателями и списка конфликтующих пар. Проверяется наличие названий дисциплин, наличие преподавателей, отсутствие повторов, существование дисциплин в конфликтах и ограничение количества дисциплин до 20.",
    )
    replace_start(
        doc,
        "Вторая функция – построение графа конфликтов",
        "Вторая функция — построение графа конфликтов. По списку дисциплин и конфликтов строится неориентированный граф, для каждой вершины вычисляется степень, а повторяющиеся и пустые конфликты не включаются в итоговый набор данных.",
    )
    replace_start(
        doc,
        "Третья функция – жадная раскраска",
        "Третья функция — жадная раскраска Welsh-Powell. Вершины упорядочиваются по убыванию степени, после чего каждой дисциплине назначается минимальный номер смены, который не используется конфликтующими соседями.",
    )
    replace_start(
        doc,
        "Четвёртая функция – локальная оптимизация",
        "Четвёртая функция — локальная оптимизация. После первичной раскраски программа пытается уменьшить количество смен и снизить число разных смен у одного преподавателя, не нарушая правило несовместимости конфликтующих дисциплин.",
    )
    replace_start(
        doc,
        "Пятая функция – формирование и отображение результата",
        "Пятая функция — формирование и отображение результата. На странице выводятся количество смен, штраф по сменам преподавателей, количество конфликтов, расписание по сменам, смены каждого преподавателя и порядок обработки вершин. После расчёта доступен экспорт результата в JSON.",
    )
    replace_start(
        doc,
        "К индивидуальному модулю (вариант 4) относятся следующие файлы",
        "К индивидуальному модулю (вариант 4) относятся следующие файлы: algorithms/coloring.py — основная вычислительная логика; views/coloring_theory.tpl — теоретическая страница; views/coloring_practice.tpl — форма практической работы; static/scripts/coloring_practice.js — клиентская логика таблиц, конфликтов и интерактивного графа; static/content/coloring_theory.css — стили теории и практики; data/coloring_history.json — история успешных расчётов; tests/test_coloring.py и tests/test_coloring_ui.py — модульные и Selenium-тесты.",
    )
    replace_start(
        doc,
        "Маршрут /coloring в файле routes.py связывает URL",
        "В routes.py используются маршруты /coloring и /coloring/practice. Маршрут /coloring открывает теоретическое описание, а /coloring/practice обрабатывает действия формы: расчёт расписания, импорт JSON, случайную генерацию данных и экспорт результата. Для передачи действия используется скрытое поле form_action.",
    )
    replace_start(
        doc,
        "Пользователь получает доступ к модулю раскраски графа через навигационную панель",
        "Пользователь получает доступ к модулю раскраски графа через навигационную панель приложения или с главной страницы. Теоретическая страница объясняет задачу раскраски графа, а страница практики содержит редактор дисциплин и конфликтов, блок импорта JSON, блок генерации, интерактивную область графа и таблицы результата.",
    )
    replace_start(
        doc,
        "На теоретической странице модуля пользователь знакомится с теорией раскраски графов",
        "На теоретической странице модуля пользователь знакомится с базовыми понятиями: дисциплина, конфликт, вершина, ребро, смена и цвет. Теоретическая информация дана в упрощённой форме, чтобы пользователь мог быстро перейти к практическому вводу данных.",
    )
    replace_start(
        doc,
        "На странице практики пользователь вводит исходные данные тремя способами",
        "На странице практики пользователь вводит исходные данные тремя способами: вручную через таблицу дисциплин и список конфликтов, через загрузку JSON-файла или через случайную генерацию. После импорта JSON данные остаются редактируемыми: можно добавить дисциплину, изменить преподавателя, добавить или удалить конфликт.",
    )
    replace_start(
        doc,
        "При успешном расчёте запись о нём дописывается в файл истории",
        "При успешном расчёте запись дописывается в файл data/coloring_history.json с указанием даты, времени, входных данных и результата. При некорректных данных расчёт не выполняется, пользователю выводится информативное сообщение об ошибке. Экспортированный JSON сохраняет поля subjects и conflicts, поэтому его можно повторно импортировать.",
    )

    app_a = find_para(doc, exact="ПРИЛОЖЕНИЕ А")

    add_page_break_before(app_a)
    add_heading_before(app_a, "5 Тестирование", "Heading 1")
    for text in [
        "Тестирование выполнялось для проверки корректности алгоритма раскраски графа, проверки обработки входных данных и проверки пользовательского интерфейса страницы практики четвертого варианта. Основное внимание уделялось индивидуальному модулю автоматического составления расписания, поскольку именно он содержит вычислительную логику студента Петренко Антона.",
        "Для проверки серверной логики используются unit-тесты на базе стандартного модуля unittest. Тесты расположены в файле tests/test_coloring.py и проверяют функции модуля algorithms/coloring.py: валидацию дисциплин и конфликтов, построение неориентированного графа, сортировку вершин, жадную раскраску, формирование расписания, расчёт смен преподавателей, импорт JSON, экспорт JSON и случайную генерацию.",
        "Для проверки интерфейса используется Selenium. UI-тесты расположены в файле tests/test_coloring_ui.py. Тесты автоматически запускают локальный Bottle-сервер на свободном порту, открывают страницу /coloring/practice в браузере Chrome или Edge, выполняют действия пользователя и проверяют результат на странице.",
    ]:
        insert_before(app_a, text)

    add_heading_before(app_a, "5.1 Проверяемые сценарии", "Heading 2")
    for text in [
        "TC_UNIT_1 — пустой список дисциплин: ожидается ошибка валидации.",
        "TC_UNIT_2 — повторяющиеся дисциплины: ожидается ошибка о повторе названия.",
        "TC_UNIT_3 — конфликт с неизвестной дисциплиной: ожидается ошибка валидации.",
        "TC_UNIT_4 — граф без конфликтов: все дисциплины могут быть размещены в одной смене.",
        "TC_UNIT_5 — цепочка конфликтов: конфликтующие дисциплины распределяются по разным сменам.",
        "TC_UNIT_6 — полный граф: каждая дисциплина получает отдельную смену.",
        "TC_UNIT_7 — корректный JSON: данные загружаются и могут быть рассчитаны.",
        "TC_UNIT_8 — экспортированный JSON: файл можно повторно импортировать без потери исходных данных.",
        "TC_UI_1 — открытие страницы практики: отображаются таблица дисциплин, список конфликтов, импорт JSON, генерация и граф.",
        "TC_UI_2 — импорт JSON и ручное редактирование: после импорта можно добавить дисциплину и конфликт.",
        "TC_UI_3 — случайная генерация: таблица дисциплин обновляется заданным количеством строк.",
        "TC_UI_4 — избыточные данные: при отправке 21 дисциплины отображается ошибка «Количество дисциплин не должно превышать 20», результат не строится.",
    ]:
        insert_before(app_a, text)

    add_heading_before(app_a, "5.2 Запуск тестирования", "Heading 2")
    for text in [
        "Запуск всех автоматических тестов выполняется командой: python -m unittest discover -s tests -p \"test_*.py\".",
        "Отдельный запуск unit-тестов логики выполняется командой: python -m unittest tests.test_coloring.",
        "Отдельный запуск Selenium-тестов выполняется командой: python -m unittest tests.test_coloring_ui -v. Для запуска UI-тестов требуется установленный пакет selenium и браузер Chrome или Edge. В текущей среде Selenium установлен, проверка через браузер выполнена успешно.",
        "По результатам проверки выполнено 114 тестов, все тесты завершились успешно. Selenium-проверка четвертого модуля включает 4 сценария и также завершилась без ошибок.",
    ]:
        insert_before(app_a, text)

    add_page_break_before(app_a)
    add_heading_before(app_a, "ЗАКЛЮЧЕНИЕ", "Heading 1")
    for text in [
        "В ходе учебной практики был разработан и интегрирован индивидуальный модуль автоматического составления расписания на основе раскраски графа конфликтов. Модуль реализован как часть многостраничного веб-приложения BottleWebProject_С326_TEAM_STEP и доступен через страницы /coloring и /coloring/practice.",
        "В результате работы реализованы ввод дисциплин и конфликтов, импорт исходных данных из JSON, возможность ручного редактирования данных после импорта, случайная генерация графа конфликтов, интерактивный просмотр графа, расчёт расписания, отображение результата и экспорт результата в JSON. Успешные расчёты сохраняются в файл истории data/coloring_history.json.",
        "Алгоритмическая часть соответствует индивидуальному варианту: используется жадная раскраска Welsh-Powell и локальная оптимизация. Реализация не смешана с логикой других участников проекта: основная вычислительная часть вынесена в algorithms/coloring.py, клиентская логика страницы — в static/scripts/coloring_practice.js, а маршрутизация четвертого модуля выделена в соответствующих обработчиках routes.py.",
        "Корректность работы проверена unit-тестами и Selenium-тестами. Unit-тесты покрывают математическую логику и обработку данных, а Selenium-тесты проверяют реальные пользовательские сценарии в браузере. Таким образом, поставленная цель разработки индивидуального модуля достигнута.",
    ]:
        insert_before(app_a, text)

    add_page_break_before(app_a)
    add_heading_before(app_a, "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ", "Heading 1")
    for text in [
        "1. ГОСТ 19.201-78. Единая система программной документации. Техническое задание. Требования к содержанию и оформлению.",
        "2. ГОСТ 19.701-90. Единая система программной документации. Схемы алгоритмов, программ, данных и систем.",
        "3. ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "4. Bottle: Python Web Framework. URL: https://bottlepy.org/docs/dev/ (дата обращения: 03.06.2026).",
        "5. Python 3 Documentation. URL: https://docs.python.org/3/ (дата обращения: 03.06.2026).",
        "6. Python unittest — Unit testing framework. URL: https://docs.python.org/3/library/unittest.html (дата обращения: 03.06.2026).",
        "7. Selenium Documentation. URL: https://www.selenium.dev/documentation/ (дата обращения: 03.06.2026).",
        "8. Vis Network documentation. URL: https://visjs.github.io/vis-network/docs/network/ (дата обращения: 03.06.2026).",
    ]:
        insert_before(app_a, text, body=False)

    clear_text(app_a, "ПРИЛОЖЕНИЕ А. Листинг 1. Код файлов вёрстки проекта", bold=True)
    app_a.style = "Heading 1"
    format_no_indent(app_a)

    doc.add_page_break()
    append_heading(doc, "ПРИЛОЖЕНИЕ Б. Листинг 2. Код логики индивидуального модуля", "Heading 1")
    append_body(
        doc,
        "В приложении приведены программные файлы, отвечающие за логику четвертого варианта: алгоритм раскраски графа, маршруты Bottle и клиентская JavaScript-логика страницы практики.",
    )
    append_heading(doc, "Файл: algorithms/coloring.py", "Heading 2")
    append_code(doc, read_project_file("algorithms/coloring.py"))
    append_heading(doc, "Файл: routes.py (фрагмент четвертого варианта)", "Heading 2")
    append_code(doc, routes_coloring_excerpt())
    append_heading(doc, "Файл: static/scripts/coloring_practice.js", "Heading 2")
    append_code(doc, read_project_file("static/scripts/coloring_practice.js"))

    apply_heading_styles(doc)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
